from docx import Document
from pathlib import Path

path = Path(r'C:\Users\TCHEDRE\Desktop\Honorine\portsentry.docx')
doc = Document(path)

print('Paragraphs:', len(doc.paragraphs))
count = 0
for i, p in enumerate(doc.paragraphs, 1):
    text = p.text.strip()
    if not text:
        continue
    style = p.style.name if p.style is not None else 'None'
    print(f'{i:03d} [{style}] {text}')
    count += 1
    if count >= 120:
        break

print('--- Tables:', len(doc.tables))
for ti, t in enumerate(doc.tables, 1):
    print('Table', ti, 'rows', len(t.rows), 'cols', len(t.columns))
    for row in t.rows[:5]:
        print(' | '.join(cell.text.replace('\n', ' ') for cell in row.cells))

print('--- Media relations ---')
rels = [r for r in doc.part.rels.values() if 'image' in r.target_ref or 'embeddings' in r.target_ref]
print('Media relations count:', len(rels))
for r in rels[:20]:
    print(' -', r.rId, r.reltype, r.target_ref)
